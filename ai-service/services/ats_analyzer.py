from __future__ import annotations

import os
import re
import hashlib
import logging
from dataclasses import dataclass
from typing import List, Tuple

from docx import Document
from pypdf import PdfReader

from models.ats import (
  ATSScanRequest,
  ATSScanResponse,
  EvidenceGap,
  Finding,
  KeywordAnalysis,
  KeywordBucket,
  MissingKeyword,
  RewriteStep,
  ScoreBlock,
  SectionFeedback,
  SkillsAnalysis,
  SynonymNote,
)
from services.rse_engine import build_requirements, calculate_scores, evaluate_requirements
from models.job import JobDescriptionRequest
from models.resume import ResumeParseRequest
from services.jd_parser import parse_job_description
from services.resume_parser import ResumeParser
from services.skill_utils import aliases_for, normalize_token, normalize_skill_list

logger = logging.getLogger(__name__)


_STOPWORDS = {
  'a', 'an', 'and', 'are', 'as', 'at', 'be', 'by', 'for', 'from', 'has', 'have', 'in', 'is', 'it',
  'of', 'on', 'or', 'our', 'that', 'the', 'their', 'this', 'to', 'we', 'with', 'you', 'your',
  'will', 'able', 'years', 'year', 'experience', 'role', 'required', 'requirements', 'must'
}


@dataclass
class _FormatStats:
  ext: str
  extracted_chars: int
  scanned_pdf_suspected: bool
  docx_tables: int
  docx_has_header_footer_text: bool
  multi_column_suspected: bool
  nonstandard_headings: List[str]


class ATSAnalyzer:
  def __init__(self) -> None:
    self._resume_parser = ResumeParser()

  def scan(self, payload: ATSScanRequest) -> ATSScanResponse:
    resume_len = len(payload.resume_text or '') if payload.resume_text is not None else 0
    jd_len = len(payload.job_description or '')
    resume_hash = hashlib.sha256((payload.resume_text or '').encode('utf-8')).hexdigest()[:10] if payload.resume_text else ''
    logger.info(
      'ats_scan_start',
      extra={
        'event': 'ats_scan_start',
        'job_id': payload.job_id,
        'resume_id': payload.resume_id,
        'resume_hash': resume_hash,
        'resume_len': resume_len,
        'jd_len': jd_len
      }
    )
    # Parse resume (reuse existing pipeline)
    resume_parse = self._resume_parser.parse(
      ResumeParseRequest(
        file_path=payload.file_path,
        file_name=payload.file_name,
        user_id=payload.user_id,
        resume_text=payload.resume_text,
        candidate_name=payload.candidate_name
      )
    )

    # Parse JD (reuse existing pipeline)
    jd_parse = parse_job_description(
      JobDescriptionRequest(
        job_title=payload.job_title,
        job_description=payload.job_description,
        location=None
      )
    )

    resume_text = (payload.resume_text or '')
    if not resume_text:
      # ResumeParser already extracted; use internal method's output by re-extracting.
      # We do a cheap extraction again here to keep code isolated.
      resume_text = self._extract_text(payload.file_path)

    jd_text = (payload.job_description or '').strip()

    # Format findings
    format_stats = self._inspect_format(payload.file_path, resume_text)
    format_findings = self._build_format_findings(format_stats)
    ats_readability_score = self._score_readability(format_findings, format_stats)

    # RSE: build & evaluate requirements once, reuse across outputs
    requirements = build_requirements(jd_text, jd_parse)
    requirement_results = evaluate_requirements(requirements, resume_text, resume_parse)
    jd_score = calculate_scores(requirements, requirement_results)

    req_index = {req.id: req for req in requirements}

    matched_skills: List[str] = []
    missing_required: List[str] = []
    missing_preferred: List[str] = []
    for res in requirement_results:
      req = req_index.get(res.requirementId)
      if not req or req.type != 'skill':
        continue
      if res.status == 'MISSING':
        if req.isRequired:
          missing_required.append(req.rawText)
        else:
          missing_preferred.append(req.rawText)
      else:
        matched_skills.append(req.rawText)

    def _bucket(req_filter):
      matched: List[str] = []
      missing: List[MissingKeyword] = []
      for res in requirement_results:
        req = req_index.get(res.requirementId)
        if not req or not req_filter(req):
          continue
        if res.status == 'MISSING':
          missing.append(
            MissingKeyword(
              keyword=req.rawText,
              importance=5 if req.isRequired else 3,
              jdEvidence=[req.rawText],
              suggestedPlacement='Experience' if req.isRequired else 'Skills'
            )
          )
        else:
          matched.append(req.rawText)
      return KeywordBucket(matched=list(dict.fromkeys(matched)), missing=missing[:25])

    kw_required_bucket = _bucket(lambda r: r.isRequired)
    kw_preferred_bucket = _bucket(lambda r: not r.isRequired)

    total_required = len(kw_required_bucket.matched) + len(kw_required_bucket.missing)
    keyword_match_score = int(round((len(kw_required_bucket.matched) / (total_required or 1)) * 100))

    evidence_gaps: List[EvidenceGap] = []
    for res in requirement_results:
      req = req_index.get(res.requirementId)
      if not req or not req.isRequired or req.type != 'skill':
        continue
      if res.status == 'WEAK':
        evidence_gaps.append(
          EvidenceGap(
            requirement=f'Show evidence for {req.rawText}',
            status='weak',
            exampleFix=(
              f'Add a bullet describing how you used {req.rawText} with outcomes. '
              'Use real work only; avoid fabrication.'
            ),
            whereToAdd='Experience'
          )
        )
      if res.status == 'MISSING':
        evidence_gaps.append(
          EvidenceGap(
            requirement=f'Missing required skill: {req.rawText}',
            status='missing',
            exampleFix=(
              f'Only add {req.rawText} if you truly have it. '
              'List it in Skills and add one supporting bullet in Experience/Projects.'
            ),
            whereToAdd='Experience'
          )
        )

    evidence_score = int(round(jd_score.evidenceStrengthScore))

    section_feedback = self._section_feedback(resume_text)

    rewrite_plan = self._rewrite_plan(format_findings, kw_required_bucket, evidence_gaps)

    synonym_notes: List[SynonymNote] = []

    response = ATSScanResponse(
      jobId=payload.job_id,
      resumeId=payload.resume_id,
      overall=ScoreBlock(
        atsReadabilityScore=ats_readability_score,
        keywordMatchScore=keyword_match_score,
        evidenceScore=evidence_score,
        jdFitScore=jd_score.jdFitScore
      ),
      jdScore=jd_score,
      requirementResults=requirement_results,
      formatFindings=format_findings,
      keywordAnalysis=KeywordAnalysis(
        required=kw_required_bucket,
        preferred=kw_preferred_bucket
      ),
      skills=SkillsAnalysis(
        matched=matched_skills,
        missingRequired=missing_required,
        missingPreferred=missing_preferred,
        synonymNotes=synonym_notes
      ),
      evidenceGaps=evidence_gaps,
      sectionFeedback=section_feedback,
      rewritePlan=rewrite_plan
    )
    logger.info(
      'ats_scan_success',
      extra={
        'event': 'ats_scan_success',
        'job_id': payload.job_id,
        'resume_id': payload.resume_id,
        'findings_count': len(format_findings),
        'missing_required': len(missing_required),
        'missing_preferred': len(missing_preferred),
        'evidence_gaps': len(evidence_gaps),
        'rse_counts': jd_score.counts,
        'scores': {
          'readability': ats_readability_score,
          'keyword': keyword_match_score,
          'evidence': evidence_score,
          'jd_fit': jd_score.jdFitScore
        }
      }
    )
    return response

  def _extract_text(self, file_path: str) -> str:
    if not file_path or not os.path.exists(file_path):
      return ''
    ext = os.path.splitext(file_path)[1].lower()
    try:
      if ext == '.pdf':
        reader = PdfReader(file_path)
        return '\n'.join((p.extract_text() or '') for p in reader.pages).strip()
      if ext == '.docx':
        doc = Document(file_path)
        return '\n'.join(p.text for p in doc.paragraphs).strip()
      with open(file_path, 'r', encoding='utf-8', errors='ignore') as handle:
        return handle.read().strip()
    except Exception:
      return ''

  def _inspect_format(self, file_path: str | None, resume_text: str) -> _FormatStats:
    ext = os.path.splitext(file_path or '')[1].lower()
    extracted = self._extract_text(file_path) if file_path else (resume_text or '')
    if not extracted and resume_text:
      extracted = resume_text
    extracted_chars = len(extracted)
    scanned_pdf_suspected = False
    docx_tables = 0
    header_footer_text = False
    multi_column_suspected = False
    nonstandard_headings: List[str] = []

    if ext == '.pdf':
      # Heuristic: if we can't extract much text, it's likely scanned or image-heavy.
      scanned_pdf_suspected = extracted_chars < 250
    elif ext == '.docx':
      try:
        doc = Document(file_path)
        docx_tables = len(doc.tables)
        # Detect any non-empty header/footer content.
        for section in doc.sections:
          header_text = ' '.join(p.text.strip() for p in section.header.paragraphs if p.text.strip())
          footer_text = ' '.join(p.text.strip() for p in section.footer.paragraphs if p.text.strip())
          if header_text or footer_text:
            header_footer_text = True
            break
      except Exception:
        pass

    lines = [l.strip() for l in extracted.splitlines() if l.strip()]
    if lines:
      column_like = sum(1 for l in lines if re.search(r'\s{2,}', l) and len(l) <= 80)
      short_lines = sum(1 for l in lines if len(l) <= 25)
      multi_column_suspected = column_like >= 5 or (short_lines >= 12 and short_lines / max(len(lines), 1) > 0.35)

      known_headings = {
        'summary', 'profile', 'about', 'skills', 'experience', 'work experience',
        'education', 'projects', 'project', 'certifications', 'achievements'
      }
      for line in lines:
        if len(line) > 60:
          continue
        candidate = line.rstrip(':').strip()
        if not candidate:
          continue
        if re.match(r'^[A-Za-z][A-Za-z0-9 .&/-]{2,}$', candidate):
          normalized = candidate.lower()
          if normalized not in known_headings and (candidate.isupper() or candidate.istitle()):
            nonstandard_headings.append(candidate)
        if len(nonstandard_headings) >= 5:
          break

    return _FormatStats(
      ext=ext,
      extracted_chars=extracted_chars,
      scanned_pdf_suspected=scanned_pdf_suspected,
      docx_tables=docx_tables,
      docx_has_header_footer_text=header_footer_text,
      multi_column_suspected=multi_column_suspected,
      nonstandard_headings=nonstandard_headings
    )

  def _build_format_findings(self, stats: _FormatStats) -> List[Finding]:
    findings: List[Finding] = []

    if stats.ext == '.pdf' and stats.scanned_pdf_suspected:
      findings.append(
        Finding(
          severity='critical',
          code='SCANNED_PDF',
          message='This PDF looks like it may be scanned or image-based (very little extractable text).',
          whyItMatters='Many ATS parsers rely on extractable text. Scanned PDFs can lead to missing or jumbled content.',
          fix='Export the resume as a text-based PDF or (best) a simple DOCX without graphics or columns.'
        )
      )

    if stats.ext == '.docx' and stats.docx_tables > 0:
      findings.append(
        Finding(
          severity='warning',
          code='TABLES_USED',
          message=f'DOCX contains {stats.docx_tables} table(s).',
          whyItMatters='Tables/columns can cause ATS systems to read content out of order or skip fields.',
          fix='Replace tables with a single-column layout using simple headings and bullet lists.'
        )
      )

    if stats.ext == '.docx' and stats.docx_has_header_footer_text:
      findings.append(
        Finding(
          severity='warning',
          code='HEADERS_FOOTERS_USED',
          message='DOCX contains text in header/footer.',
          whyItMatters='Some ATS parsers ignore header/footer text, potentially losing contact details or section content.',
          fix='Move critical information (name, email, phone, links) into the main document body.'
        )
      )

    if stats.extracted_chars < 150 and stats.ext not in {'.pdf', '.docx'}:
      findings.append(
        Finding(
          severity='warning',
          code='LOW_TEXT_EXTRACT',
          message='Low amount of extractable text detected.',
          whyItMatters='If the system cannot read text reliably, matching will be inaccurate.',
          fix='Upload a text-based PDF or DOCX. Avoid scanned images.'
        )
      )

    if stats.ext not in {'.pdf', '.docx', '.txt'}:
      findings.append(
        Finding(
          severity='info',
          code='UNKNOWN_FORMAT',
          message=f'File extension {stats.ext or "(none)"} is not explicitly optimized.',
          whyItMatters='Some ATS systems handle DOCX best and may struggle with uncommon formats.',
          fix='Prefer DOCX or a simple, text-based PDF.'
        )
      )

    if stats.multi_column_suspected:
      findings.append(
        Finding(
          severity='warning',
          code='MULTI_COLUMN_LAYOUT',
          message='Resume text layout looks multi-column or table-like.',
          whyItMatters='Multi-column layouts can confuse ATS parsers and jumble content.',
          fix='Use a single-column layout with simple headings and bullet lists.'
        )
      )

    if stats.nonstandard_headings:
      sample = ', '.join(stats.nonstandard_headings[:4])
      findings.append(
        Finding(
          severity='info',
          code='NONSTANDARD_HEADINGS',
          message='Detected uncommon section headings.',
          whyItMatters='ATS systems prefer standard headings to reliably parse sections.',
          fix=f'Consider using standard headings like Skills, Experience, Education. Found: {sample}'
        )
      )

    return findings

  def _score_readability(self, findings: List[Finding], stats: _FormatStats) -> int:
    score = 100
    for f in findings:
      if f.code == 'SCANNED_PDF':
        score -= 50
      elif f.code == 'TABLES_USED':
        score -= 25
      elif f.code == 'HEADERS_FOOTERS_USED':
        score -= 10
      elif f.code == 'LOW_TEXT_EXTRACT':
        score -= 20
      elif f.code == 'MULTI_COLUMN_LAYOUT':
        score -= 15
      elif f.code == 'NONSTANDARD_HEADINGS':
        score -= 5
    if stats.extracted_chars < 500:
      score -= 5
    return max(0, min(100, score))

  def _compare_skills(
    self,
    resume_skills: List[str],
    required: List[str],
    preferred: List[str],
    resume_text: str
  ) -> Tuple[List[str], List[str], List[str], List[SynonymNote]]:
    resume_norm = {s.lower(): s for s in resume_skills}
    text_norm = normalize_token(resume_text)

    matched: List[str] = []
    missing_required: List[str] = []
    missing_preferred: List[str] = []
    synonym_notes: List[SynonymNote] = []

    def _has_skill(skill: str) -> bool:
      key = skill.lower()
      if key in resume_norm:
        return True
      # alias check (skill aliases in text)
      for alias in aliases_for(skill):
        if alias in text_norm:
          synonym_notes.append(SynonymNote(resumeTerm=alias, jdTerm=skill, treatedAsMatch=True))
          return True
      return False

    for skill in required:
      if _has_skill(skill):
        matched.append(skill)
      else:
        missing_required.append(skill)

    for skill in preferred:
      if _has_skill(skill):
        if skill not in matched:
          matched.append(skill)
      else:
        missing_preferred.append(skill)

    return matched, missing_required, missing_preferred, synonym_notes

  def _extract_keywords(self, jd_text: str) -> Tuple[List[str], List[str]]:
    """Return (required_keywords, preferred_keywords) from job text.

    This is intentionally light-weight (no external NLP deps).
    """
    lines = [l.strip() for l in jd_text.splitlines() if l.strip()]

    required_block: List[str] = []
    preferred_block: List[str] = []

    preferred_hints = ('nice to have', 'preferred', 'bonus', 'plus', 'optional')

    current = 'required'
    for line in lines:
      lower = line.lower()
      if any(h in lower for h in preferred_hints):
        current = 'preferred'
      if current == 'preferred':
        preferred_block.append(line)
      else:
        required_block.append(line)

    def _tokens(block_lines: List[str]) -> List[str]:
      text = ' '.join(block_lines)
      # pick multi-word phrases like "rest api", "micro services" as well as single tokens
      raw_tokens = re.findall(r"[A-Za-z][A-Za-z0-9+./#-]{1,}", text)
      cleaned: List[str] = []
      for tok in raw_tokens:
        t = tok.strip('-').strip()
        if len(t) < 3:
          continue
        if t.lower() in _STOPWORDS:
          continue
        cleaned.append(t)
      # frequency sort
      freq = {}
      for t in cleaned:
        key = t.lower()
        freq[key] = freq.get(key, 0) + 1
      ordered = sorted(freq.items(), key=lambda kv: (-kv[1], kv[0]))
      # keep top N
      return [t for t, _ in ordered[:35]]

    return _tokens(required_block), _tokens(preferred_block)

  def _merge_keywords(self, keywords: List[str], skills: List[str]) -> List[str]:
    merged = {k.lower(): k for k in keywords}
    for s in skills:
      merged.setdefault(s.lower(), s)
    return list(merged.values())

  def _bucket_keywords(self, keywords: List[str], resume_text: str) -> KeywordBucket:
    text_norm = normalize_token(resume_text)
    matched: List[str] = []
    missing: List[MissingKeyword] = []

    for kw in keywords:
      key = normalize_token(kw)
      if not key or key in _STOPWORDS:
        continue
      if key in text_norm:
        matched.append(kw)
      else:
        missing.append(
          MissingKeyword(
            keyword=kw,
            importance=3,
            jdEvidence=[kw],
            suggestedPlacement='Skills'
          )
        )

    # dedupe / stable
    matched = list(dict.fromkeys(matched))
    missing = missing[:25]
    return KeywordBucket(matched=matched, missing=missing)

  def _score_keyword_match(self, required_bucket: KeywordBucket, preferred_bucket: KeywordBucket) -> int:
    req_total = len(required_bucket.matched) + len(required_bucket.missing)
    pref_total = len(preferred_bucket.matched) + len(preferred_bucket.missing)

    def _ratio(m: int, total: int) -> float:
      return 1.0 if total <= 0 else m / total

    req_ratio = _ratio(len(required_bucket.matched), req_total)
    pref_ratio = _ratio(len(preferred_bucket.matched), pref_total)

    score = int(round((req_ratio * 0.75 + pref_ratio * 0.25) * 100))
    return max(0, min(100, score))

  def _check_evidence(self, jd_text: str, resume_text: str, required_skills: List[str]) -> Tuple[List[EvidenceGap], int]:
    """Simple evidence heuristic.

    - 'ok' if skill appears in experience/project-ish lines
    - 'weak' if only appears somewhere else
    - 'missing' if not present
    """
    lines = [l.strip() for l in resume_text.splitlines() if l.strip()]
    exp_lines = [l for l in lines if re.search(r'\b(led|built|developed|designed|implemented|owned|migrated|deployed)\b', l, re.I)]
    exp_text = normalize_token(' '.join(exp_lines))
    full_text = normalize_token(resume_text)

    gaps: List[EvidenceGap] = []

    ok_count = 0
    for skill in required_skills:
      skill_norm = normalize_token(skill)
      if skill_norm and skill_norm in exp_text:
        ok_count += 1
        continue
      if skill_norm and skill_norm in full_text:
        gaps.append(
          EvidenceGap(
            requirement=f'Show evidence for {skill}',
            status='weak',
            exampleFix=(
              f'Add a bullet under your most relevant role describing how you used {skill} '
              'to deliver a measurable outcome (no new claims—use your real work).'
            ),
            whereToAdd='Experience'
          )
        )
      else:
        gaps.append(
          EvidenceGap(
            requirement=f'Missing required skill: {skill}',
            status='missing',
            exampleFix=(
              f'Only add {skill} if you truly have experience with it. If you do, '
              'mention it in Skills and add one supporting bullet in Experience/Projects.'
            ),
            whereToAdd='Experience'
          )
        )

    total = len(required_skills) or 1
    evidence_ratio = ok_count / total
    score = int(round(evidence_ratio * 100))
    # weak gaps penalize
    score = max(0, min(100, score - (5 * sum(1 for g in gaps if g.status == 'weak'))))
    return gaps[:20], score

  def _section_feedback(self, resume_text: str) -> List[SectionFeedback]:
    text = resume_text.lower()
    feedback: List[SectionFeedback] = []
    sections = {
      'Skills': bool(re.search(r'^\s*skills\b', text, re.M)),
      'Experience': bool(re.search(r'^\s*(experience|work experience)\b', text, re.M)),
      'Education': bool(re.search(r'^\s*education\b', text, re.M)),
      'Projects': bool(re.search(r'^\s*projects?\b', text, re.M)),
      'Summary': bool(re.search(r'^\s*(summary|profile|about)\b', text, re.M)),
    }

    if not sections['Skills']:
      feedback.append(
        SectionFeedback(
          section='Skills',
          severity='warning',
          message='No clear "Skills" section heading detected.',
          fix='Add a "Skills" heading and list tools/technologies as simple comma-separated items.'
        )
      )

    if not sections['Experience']:
      feedback.append(
        SectionFeedback(
          section='Experience',
          severity='warning',
          message='No clear "Experience"/"Work Experience" heading detected.',
          fix='Use a standard "Work Experience" heading with roles in reverse-chronological order.'
        )
      )

    return feedback

  def _rewrite_plan(
    self,
    format_findings: List[Finding],
    kw_required: KeywordBucket,
    evidence_gaps: List[EvidenceGap]
  ) -> List[RewriteStep]:
    plan: List[RewriteStep] = []

    for f in format_findings:
      if f.severity == 'critical':
        plan.append(
          RewriteStep(
            priority='P0',
            title='Fix resume format for ATS readability',
            action=f.fix,
            details=f.message
          )
        )

    if kw_required.missing:
      top_missing = ', '.join(m.keyword for m in kw_required.missing[:8])
      plan.append(
        RewriteStep(
          priority='P0',
          title='Close critical keyword gaps',
          action='Update Skills section (and add evidence bullets where applicable).',
          details=f'Missing keywords: {top_missing}'
        )
      )

    weak_evidence = [g for g in evidence_gaps if g.status == 'weak']
    if weak_evidence:
      plan.append(
        RewriteStep(
          priority='P1',
          title='Strengthen evidence in Experience bullets',
          action='Add 1–2 bullets showing real usage of required skills.',
          details='Focus on results (latency, throughput, costs, reliability) and your actual contributions.'
        )
      )

    return plan[:10]


def ats_scan(payload: ATSScanRequest) -> ATSScanResponse:
  return ATSAnalyzer().scan(payload)
